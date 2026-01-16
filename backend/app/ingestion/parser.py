from pathlib import Path
from typing import List, Dict
import re
import uuid

import pdfplumber
from docx import Document

from .table_extractor import extract_tables


# -----------------------------
# Utilities
# -----------------------------

def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def _split_paragraphs(text: str) -> List[str]:
    return [p.strip() for p in text.split("\n") if p.strip()]


def _new_block_id() -> str:
    return uuid.uuid4().hex


# -----------------------------
# PDF Parsing
# -----------------------------

def parse_pdf(file_path: Path) -> List[Dict]:
    blocks: List[Dict] = []

    # -------- TEXT EXTRACTION --------
    with pdfplumber.open(file_path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_number = page_index + 1

            raw_text = page.extract_text()
            if not raw_text:
                continue

            cleaned = _clean_text(raw_text)
            paragraphs = _split_paragraphs(cleaned)

            for para_index, paragraph in enumerate(paragraphs):
                blocks.append({
                    "block_id": _new_block_id(),
                    "text": paragraph,
                    "source": file_path.name,
                    "page": page_number,
                    "doc_level": page_index == 0 and para_index == 0,
                    "block_type": "text",
                })

    # -------- TABLE EXTRACTION --------
    table_rows = extract_tables(file_path)

    for row in table_rows:
        blocks.append({
            "block_id": row["block_id"],
            "text": row["text"],
            "source": file_path.name,
            "page": row.get("page"),
            "doc_level": False,
            "block_type": "table_row",
            "table_id": row.get("table_id"),
            "row_index": row.get("table_row"),
        })

    return blocks


# -----------------------------
# DOCX Parsing
# -----------------------------

def parse_docx(file_path: Path) -> List[Dict]:
    doc = Document(file_path)
    blocks: List[Dict] = []

    # -------- PARAGRAPHS --------
    for para_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        blocks.append({
            "block_id": _new_block_id(),
            "text": _clean_text(text),
            "source": file_path.name,
            "page": None,
            "doc_level": para_index == 0,
            "block_type": "text",
        })

    # -------- TABLES --------
    for table_index, table in enumerate(doc.tables):
        table_id = f"docx_table_{table_index}"

        for row_index, row in enumerate(table.rows):
            values = [cell.text.strip() for cell in row.cells]
            if not any(values):
                continue

            blocks.append({
                "block_id": _new_block_id(),
                "text": " | ".join(values),
                "source": file_path.name,
                "page": None,
                "doc_level": False,
                "block_type": "table_row",
                "table_id": table_id,
                "row_index": row_index,
            })

    return blocks


# -----------------------------
# TXT Parsing
# -----------------------------

def parse_txt(file_path: Path) -> List[Dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        raw = f.read()

    cleaned = _clean_text(raw)
    paragraphs = _split_paragraphs(cleaned)

    blocks: List[Dict] = []

    for index, paragraph in enumerate(paragraphs):
        blocks.append({
            "block_id": _new_block_id(),
            "text": paragraph,
            "source": file_path.name,
            "page": None,
            "doc_level": index == 0,
            "block_type": "text",
        })

    return blocks



